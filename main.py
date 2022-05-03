# -*- coding: utf-8 -*-

from pydoc import doc
from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor,Pt
from datetime import date, datetime
import moment



# ENTETE = """Je soussigne Docteur d\’Etat de medecine PAPA BIRANE BOYE, certifie avoir consulte dans le cadre de le visite medicale annuelle 2021 """
ENTETE = """DOCTEUR PAPA BIRANE BOYE
N° 0499 DE L’ORDRE NATIONAL DES MEDECINS DU SENEGAL
MEDECINE DU TRAVAIL
MEDECINE GENERALE
Rue des CHAIS BEL AIR
TEL 774503252
E- Mail : pbboye@hotmail.com
BP 3750"""

BODY = """Je soussigné Docteur d’Etat de médecine PAPA BIRANE BOYE, certifie avoir consulté dans le cadre de la visite médicale annuelle 2021
"""

APTE="""Par conséquent certifions qu’il n’est atteint d’aucune maladie cliniquement ni radiologiquement décelable ;
"""

FIRSTNAME = 1
LASTNAME = 2
BIRTHDAY = 3
GENDER = 4
COMPANY = 5

wb = load_workbook('main.xlsx',{'default_date_format':'dd/mm/yy'})

ws = wb.active

print(ws.max_row)

def describe_user(firstName,lastName,birthday,gender,company):
    print('\n--------------------------')
    print('prenom : {0} \nnom : {1} \ndate de naissance : {2} \nsexe : {3} \nentreprise : {4}'.format(firstName,lastName,birthday,gender,company))

def write_document(document,prenom,nom,birthday,gender):
    
    p_entete = document.add_paragraph(ENTETE)
    p_entete.add_run('\n')
    p_entete.add_run().font.color.rgb = RGBColor.from_string('0b0fff')

    oleosen_par = document.add_paragraph('\n')

    oleosen = oleosen_par.add_run('\tOLEOSEN')
    oleosen.bold = True
    oleosen.font.color.rgb = RGBColor.from_string('993300')
    oleosen.font.size = Pt(22)
    oleosen_par.add_run('\t\t\t\t\tDakar le :')
    oleosen_par.add_run(date.today().strftime("%d/%m/%Y"))


    par = document.add_heading('',level=1)
    par.alignment =  WD_ALIGN_PARAGRAPH.CENTER
    par.add_run('\nCERTIFICAT MEDICAL\n').bold = True

    p = document.add_paragraph('\n')
    p.add_run(BODY)
    gender_title =''
    gender_title = 'Mme. ' if gender == 'Female' else 'Mr. ' 
    p.add_run(gender_title)
    p.add_run(prenom+' '+nom).bold= True
    p.add_run('\nNé (e) le :')
    
    p.add_run(birthday).bold= True
    p.add_run('\n')

    p = document.add_paragraph('\n')

    document.add_paragraph('Examen clinique normal', style='List Bullet')
    document.add_paragraph('Examen biologique normal', style='List Bullet')
    document.add_paragraph('Examen radiologique pulmonaire normal\n', style='List Bullet')


    p = document.add_paragraph('\n')

    p.add_run(APTE)
    p.add_run('En conclusion l’estimons ')
    p.add_run('APTE ').bold= True
    p.add_run('pour le travail.')
    p.add_run('\n')
    p.add_run('\n')

    par = document.add_paragraph('')
    par.alignment =  WD_ALIGN_PARAGRAPH.CENTER
    par.add_run('\nLe médecin d’entreprise').underline = True

    # new_section = document.add_section()
    document.add_page_break()


document = Document()

for row in range(1,ws.max_row+1):
    firstName = ws.cell(row,FIRSTNAME).value
    lastName = ws.cell(row,LASTNAME).value
    birthday = ws.cell(row,BIRTHDAY).value
    # birthday = datetime.strptime(ws.cell(row,BIRTHDAY).value,'%y-%m-%d %H:%M:%s')
    birthday = ws.cell(row,BIRTHDAY).value
    gender = ws.cell(row,GENDER).value
    company = ws.cell(row,COMPANY).value

    birthday = birthday.split(' ')[0]
    print(birthday)


    # Description de l'utilisateur
    # describe_user(firstName,lastName,birthday,gender,company)

    write_document(document,firstName,lastName,birthday,gender)


document.save('test.docx')




wb.close()

